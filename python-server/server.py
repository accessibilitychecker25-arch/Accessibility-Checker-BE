# server.py
import io
import os
import re
import time
import uuid
import shutil
import zipfile
import tempfile
from pathlib import Path
from typing import Dict, Any, Optional

from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse

from starlette.requests import Request
from io import BytesIO
from zipfile import ZipFile, ZIP_DEFLATED

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from starlette.background import BackgroundTask


# ---------- CONFIG ----------
PUBLIC_BASE_URL = os.environ.get("PUBLIC_BASE_URL", "http://localhost:3000")
DOWNLOAD_TTL_SEC = 15 * 60
DOWNLOAD_DIR = Path(tempfile.gettempdir()) / "docx-remediations"
DOWNLOAD_DIR.mkdir(parents=True, exist_ok=True)

# ---------- APP ----------
app = FastAPI()
# Configure CORS: make allowed origins configurable via ALLOWED_ORIGINS env var.
# Provide sensible defaults for local dev and the GitHub Pages + Vercel hosts used by the frontend.
default_origins = [
    "http://localhost:4200",
    "https://accessibilitychecker25-arch.github.io",
    "https://accessibility-checker-be.vercel.app",
]
allowed_origins_env = os.environ.get("ALLOWED_ORIGINS")
if allowed_origins_env:
    # comma-separated list in env var
    allowed_origins = [o.strip() for o in allowed_origins_env.split(",") if o.strip()]
else:
    allowed_origins = default_origins

# Allow all common methods/headers for the API; tighten in production if needed.
app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_methods=["*"],
    allow_headers=["*"],
    allow_credentials=False,
)

@app.middleware("http")
async def access_log(request: Request, call_next):
    t0 = time.time()
    resp = await call_next(request)
    dt = (time.time() - t0) * 1000
    print(f"[{request.method}] {request.url.path} -> {resp.status_code} ({dt:.1f} ms)")
    return resp

@app.get("/")
def health():
    return {"ok": True, "service": "docx-remediation"}

# ---------- UTILS ----------
def is_docx(filename: str, mime: Optional[str]) -> bool:
    return (filename or "").lower().endswith(".docx") or (
        mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

def file_name_has_problems(name: str) -> bool:
    """
    This function checks if a filename is undescriptive, based on certain patterns.
    It checks for filenames that are like 'document#', 'untitled#', or similar.
    """
    base = re.sub(r"\.docx$", "", name, flags=re.I)
    print(f"Checking if file is undescriptive: {base}")  # Debug line

    return (
        re.fullmatch(r"document\d*", base, flags=re.I) is not None
        or re.fullmatch(r"untitled\d*", base, flags=re.I) is not None
        or len(base) < 4  # If the base filename is too short (e.g., "1234.docx")
        or "_" in base  # Check if the filename contains underscores
    )

def slugify(s: str) -> str:
    """
    This function transforms a string into a slug by:
    - Removing unwanted characters
    - Replacing underscores with hyphens
    - Keeping hyphens intact
    """
    print(f"Original filename for slugify: {s}")  # Debug line
    s = re.sub(r"[^\w\s-]", "", s).strip().lower()
    s = re.sub(r"[\s_]+", "-", s)  # Replace spaces and underscores with hyphens
    s = re.sub(r"-{2,}", "-", s)  # Replace multiple hyphens with a single one
    print(f"Slugified filename: {s}")  # Debug line
    return s or "document"

def hex_to_srgb(h: str):
    h = h.strip().lstrip("#")
    r = int(h[0:2], 16) / 255.0
    g = int(h[2:4], 16) / 255.0
    b = int(h[4:6], 16) / 255.0
    def to_lin(c): return c / 12.92 if c <= 0.04045 else ((c + 0.055) / 1.055) ** 2.4
    return (to_lin(r), to_lin(g), to_lin(b))

def contrast_ratio(fg_hex: str, bg_hex: str = "FFFFFF") -> float:
    r1, g1, b1 = hex_to_srgb(fg_hex)
    r2, g2, b2 = hex_to_srgb(bg_hex)
    L1 = 0.2126*r1 + 0.7152*g1 + 0.0722*b1
    L2 = 0.2126*r2 + 0.7152*g2 + 0.0722*b2
    hi, lo = (L1, L2) if L1 >= L2 else (L2, L1)
    return (hi + 0.05) / (lo + 0.05)

def now_ts():
    return int(time.time())

def write_pkg_xml(original_bytes: bytes, replacements: Dict[str, bytes]) -> bytes:
    """
    Build a NEW .docx zip from original_bytes, replacing parts in `replacements`.
    Keys are zip member names (e.g., 'word/styles.xml'), values are raw bytes.
    """
    out = BytesIO()
    with ZipFile(BytesIO(original_bytes), "r") as zin, ZipFile(out, "w", ZIP_DEFLATED) as zout:
        replaced = set(replacements.keys())
        # write replacements first
        for name, content in replacements.items():
            zout.writestr(name, content)
        # copy everything else
        for info in zin.infolist():
            if info.filename in replaced:
                continue
            with zin.open(info.filename) as src:
                zout.writestr(info, src.read())
    return out.getvalue()

def read_xml_part(data: bytes, name: str) -> Optional[bytes]:
    try:
        with ZipFile(BytesIO(data), "r") as z:
            return z.read(name)
    except KeyError:
        return None

# ---------- LOW-RISK REMEDIATIONS ----------
def remove_protection_bytes(orig_xml: bytes) -> Optional[bytes]:
    print(orig_xml)
    root = etree.fromstring(orig_xml)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    # Remove documentProtection, writeProtection, readOnlyRecommended, editRestrictions, formProtection if present
    removed = False
    for tag in ("documentProtection", "writeProtection", "readOnlyRecommended", "editRestrictions", "formProtection"):
        el = root.find(f"w:{tag}", ns)
        if el is not None:
            el.getparent().remove(el)
            removed = True

    # Remove <w:locked/> elements and any w:locked attributes on child elements
    for locked_el in root.findall('.//w:locked', ns):
        parent = locked_el.getparent()
        if parent is not None:
            parent.remove(locked_el)
            removed = True

    # Remove w:locked attributes from any elements
    for el in root.findall('.//*'):
        if el.get(qn('w:locked')) is not None:
            try:
                del el.attrib[qn('w:locked')]
                removed = True
            except Exception:
                pass

    if not removed:
        return None
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")

def set_default_lang_en_us_bytes(orig_xml: bytes) -> Optional[bytes]:
    root = etree.fromstring(orig_xml)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    styles = root
    docDefaults = styles.find("w:docDefaults", ns)
    if docDefaults is None:
        docDefaults = etree.SubElement(styles, qn("w:docDefaults"))
    rPrDefault = docDefaults.find("w:rPrDefault", ns)
    if rPrDefault is None:
        rPrDefault = etree.SubElement(docDefaults, qn("w:rPrDefault"))
    rPr = rPrDefault.find("w:rPr", ns)
    if rPr is None:
        rPr = etree.SubElement(rPrDefault, qn("w:rPr"))
    lang = rPr.find("w:lang", ns)
    changed = False
    if lang is None:
        lang = etree.SubElement(rPr, qn("w:lang"))
        lang.set(qn("w:val"), "en-US")
        changed = True
    elif lang.get(qn("w:val")) != "en-US":
        lang.set(qn("w:val"), "en-US")
        changed = True
    if not changed:
        return None
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")

def first_heading_text(doc: Document) -> str:
    for p in doc.paragraphs:
        style = (p.style.name if p.style else "") or ""
        if re.match(r"Heading\s*[1-9]$", style, flags=re.I):
            t = p.text.strip()
            if t:
                return t
    return ""
def ensure_title_bytes(core_xml: bytes) -> Optional[bytes]:
    # Proceed with existing logic
    root = etree.fromstring(core_xml)
    ns = {
        "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
        "dc": "http://purl.org/dc/elements/1.1/",
    }
    
    # Find the <dc:title> element
    title_el = root.find("dc:title", ns)
    cur = (title_el.text or "").strip() if title_el is not None else ""

    # If the current title is undescriptive or needs to be changed, set it to "Needs Title"
    if not cur or re.match(r"(document\d*|untitled|needs title)$", cur, flags=re.I):
        # If the title is already "Needs Title", do nothing
        if title_el is None:
            title_el = etree.SubElement(root, "{%s}title" % ns["dc"])
        title_el.text = "Needs Title"  # Set it to "Needs Title"
    else:
        # If the title is already descriptive, don't change it
        return None

    # Return the modified XML
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")

def set_table_header_repeat(doc: Document, report: Dict[str, Any]):
    count = 0
    for t_index, table in enumerate(doc.tables):
        if not table.rows:
            continue
        first = table.rows[0]._tr
        trPr = first.get_or_add_trPr()  # avoids future warning
        hdr = trPr.find(qn("w:tblHeader"))
        if hdr is None:
            hdr = OxmlElement("w:tblHeader")
            hdr.set(qn("w:val"), "1")
            trPr.append(hdr)
            report["details"]["tablesHeaderRowSet"].append({"tableIndex": t_index})
            count += 1
    if count:
        report["summary"]["fixed"] += count

# ---------- DETECTION HELPERS (read-only) ----------
def detect_empty_headings_and_order(doc: Document, report: Dict[str, Any]):
    empty = []
    order = []
    prev = None
    for idx, p in enumerate(doc.paragraphs):
        style = (p.style.name if p.style else "") or ""
        m = re.match(r"Heading\s*([1-9])$", style, flags=re.I)
        if m:
            lvl = int(m.group(1))
            if not p.text.strip():
                empty.append({"paragraphIndex": idx})
            if prev is not None and lvl > prev + 1:
                order.append({
                    "paragraphIndex": idx,
                    "previousLevel": prev,
                    "currentLevel": lvl
                })
            prev = lvl
    report["details"]["emptyHeadings"] = empty
    report["details"]["headingOrderIssues"] = order
    report["summary"]["flagged"] += len(empty) + len(order)

def read_pkg_xml(zf: ZipFile, name: str) -> Optional[etree._Element]:
    try:
        with zf.open(name) as f:
            return etree.fromstring(f.read())
    except KeyError:
        return None

def detect_links(zf: ZipFile, report: Dict[str, Any]):
    doc = read_pkg_xml(zf, "word/document.xml")
    rels = read_pkg_xml(zf, "word/_rels/document.xml.rels")
    if doc is None or rels is None:
        return
    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    }
    target_by_id = {}
    for rel in rels.findall("Relationship"):
        rid = rel.get("Id")
        tgt = rel.get("Target", "")
        target_by_id[rid] = tgt

    bad = []
    for idx, p in enumerate(doc.findall(".//w:p", ns)):
        for h in p.findall(".//w:hyperlink", ns):
            rid = h.get(qn("r:id"))
            target = target_by_id.get(rid, "")
            display_parts = [t.text or "" for t in h.findall(".//w:t", ns)]
            display = "".join(display_parts).strip()
            looks_raw = (display and target and display == target)
            generic = re.search(r"\b(click here|read more|here|more)\b", display, flags=re.I) is not None
            if generic or looks_raw or len(display) > 120:
                bad.append({"paragraphIndex": idx, "display": display, "target": target or None})
    report["details"]["badLinks"] = bad
    report["summary"]["flagged"] += len(bad)

def detect_tables_merged_empty(zf: ZipFile, report: Dict[str, Any]):
    doc = read_pkg_xml(zf, "word/document.xml")
    if doc is None:
        return
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    issues = []
    for ti, tbl in enumerate(doc.findall(".//w:tbl", ns)):
        rows = tbl.findall("./w:tr", ns)
        for ri, tr in enumerate(rows):
            cells = tr.findall("./w:tc", ns)
            for ci, tc in enumerate(cells):
                tcPr = tc.find("./w:tcPr", ns)
                span = tcPr.find("./w:gridSpan", ns).get(qn("w:val")) if (tcPr is not None and tcPr.find("./w:gridSpan", ns) is not None) else None
                vMerge_el = tcPr.find("./w:vMerge", ns) if (tcPr is not None) else None
                vMerge = vMerge_el.get(qn("w:val")) if (vMerge_el is not None and vMerge_el.get(qn("w:val")) is not None) else ("" if vMerge_el is not None else None)
                textbits = [t.text or "" for t in tc.findall(".//w:t", ns)]
                text = "".join(textbits).strip()
                if span or vMerge is not None or text == "":
                    issues.append({
                        "tableIndex": ti,
                        "row": ri,
                        "col": ci,
                        "gridSpan": span or None,
                        "vMerge": vMerge,
                        "isEmpty": text == "",
                    })
    report["details"]["mergedSplitEmptyCells"] = issues
    report["summary"]["flagged"] += len(issues)

def detect_header_footer(zf: ZipFile, report: Dict[str, Any]):
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    notes = []
    for name in zf.namelist():
        if not re.match(r"word/(header|footer)\d*\.xml$", name):
            continue
        root = read_pkg_xml(zf, name)
        if root is None:
            continue
        textbits = [t.text or "" for t in root.findall(".//w:t", ns)]
        text = " ".join(textbits).strip()
        if text and len(text) >= 3:
            notes.append({
                "part": name.replace("word/", ""),
                "preview": (text[:140] + "…") if len(text) > 140 else text
            })
    report["details"]["headerFooterAudit"] = notes
    report["summary"]["flagged"] += len(notes)

def _rgb_to_hex(rgb) -> str:
    try:
        r, g, b = int(rgb[0]), int(rgb[1]), int(rgb[2])
        return f"{r:02X}{g:02X}{b:02X}"
    except Exception:
        s = str(rgb).strip().lstrip("#").upper()
        return s if re.fullmatch(r"[0-9A-F]{6}", s) else "000000"

def detect_contrast(doc: Document, report: Dict[str, Any]):
    issues = []
    for idx, p in enumerate(doc.paragraphs):
        for r in p.runs:
            color = r.font.color
            if color is None or color.rgb is None:
                continue
            hexcolor = _rgb_to_hex(color.rgb)
            ratio = contrast_ratio(hexcolor, "FFFFFF")
            size_pt = r.font.size.pt if r.font.size else None
            bold = bool(r.bold)
            required = 3.0 if (bold or (size_pt and size_pt >= 18.0)) else 4.5
            if ratio < required:
                issues.append({
                    "paragraphIndex": idx,
                    "color": hexcolor,
                    "sizePt": size_pt,
                    "bold": bold,
                    "ratio": round(ratio, 2),
                    "sample": r.text[:60],
                })
    report["details"]["colorContrastIssues"] = issues
    report["summary"]["flagged"] += len(issues)

def file_name_has_underscores(name: str) -> bool:
    """
    This function checks if a filename contains underscores.
    """
    return "_" in name

def file_name_is_untitled_or_document_followed_by_number(name: str) -> bool:
    """
    This function checks if a filename matches the pattern of 'untitled' or 'document',
    optionally followed by a number.
    """
    base = re.sub(r"\.docx$", "", name, flags=re.I)  # Remove .docx extension
    
    # Match 'untitled' or 'document' optionally followed by digits
    return bool(re.match(r"^(untitled|document)(\d*)$", base, flags=re.I))


def process_file_name(file, report):
    # **Filename suggestion and renaming logic**
    if file_name_has_problems(file.filename):  # If the filename is undescriptive (document#, untitled#)
        if file_name_has_underscores(file.filename):  # If the filename has underscores
            base = re.sub(r"\.docx$", "", file.filename, flags=re.I)
            base = base.replace("_", "-")  # Replace underscores with hyphens
            # Use the slugify function to format the base filename
            report["suggestedFileName"] = f"{slugify(base)}.docx"
            report["details"]["fileNameFixed"] = True
            report["summary"]["fixed"] += 1
        elif file_name_is_untitled_or_document_followed_by_number(file.filename):
            # If the filename is something like "untitled123" or "document123"
            report["details"]["fileNameNeedsFixing"] = True
            report["summary"]["flagged"] += 1
    else:
        # If the file name is already fine, retain it
        report["suggestedFileName"] = file.filename


# ---------- MAIN ROUTES ----------
@app.post("/upload-document")
async def upload_document(file: UploadFile = File(...), title: str = Form(default="")):

    if not file:
        raise HTTPException(400, "No file uploaded")
    if not is_docx(file.filename, file.content_type):
        raise HTTPException(400, detail={
            "error": "Please upload a .docx file",
            "details": {"received": {"name": file.filename, "mimetype": file.content_type}},
        })

    report = {
        "fileName": file.filename,
        "suggestedFileName": None,  # Initialize the suggestedFileName
        "summary": {"fixed": 0, "flagged": 0},
        "details": {
            "removedProtection": False,
            "fileNameFixed": False,
            "fileNameNeedsFixing": False,
            "titleNeedsFixing": False,
            "tablesHeaderRowSet": [],
            "emptyHeadings": [],
            "headingOrderIssues": [],
            "mergedSplitEmptyCells": [],
            "badLinks": [],
            "headerFooterAudit": [],
            "imagesMissingOrBadAlt": 0,
            "anchoredDrawingsDetected": 0,
            "embeddedMedia": [],
            "gifsDetected": [],
            "colorContrastIssues": [],
            "languageDefaultFixed": None,
        },
    }

    original_bytes = await file.read()

    # -------- Phase A: python-docx conservative edit (repeat header) --------
    tmp_path = DOWNLOAD_DIR / f"work-{uuid.uuid4().hex}.docx"
    tmp_path.write_bytes(original_bytes)
    doc = Document(str(tmp_path))
    set_table_header_repeat(doc, report)
    doc.save(str(tmp_path))
    phase_a_bytes = tmp_path.read_bytes()

    # -------- Phase B: XML replacements by rebuilding the zip --------
    replacements: Dict[str, bytes] = {}

    settings_xml = read_xml_part(phase_a_bytes, "word/settings.xml")
    if settings_xml:
        new_settings = remove_protection_bytes(settings_xml)
        if new_settings is not None:
            replacements["word/settings.xml"] = new_settings
            report["details"]["removedProtection"] = True
            report["summary"]["fixed"] += 1

    styles_xml = read_xml_part(phase_a_bytes, "word/styles.xml")
    if styles_xml:
        new_styles = set_default_lang_en_us_bytes(styles_xml)
        if new_styles is not None:
            replacements["word/styles.xml"] = new_styles
            report["details"]["languageDefaultFixed"] = {"setTo": "en-US"}
            report["summary"]["fixed"] += 1

    core_xml = read_xml_part(phase_a_bytes, "docProps/core.xml")
    if core_xml:
        new_core = ensure_title_bytes(core_xml)
        if new_core is not None:
            replacements["docProps/core.xml"] = new_core
            report["details"]["titleNeedsFixing"] = True
            report["summary"]["flagged"] += 1

    final_bytes = write_pkg_xml(phase_a_bytes, replacements)
    tmp_path.unlink(missing_ok=True)

    # -------- Phase C: detections (fresh read-only views) --------
    detect_tmp = DOWNLOAD_DIR / f"detect-{uuid.uuid4().hex}.docx"
    detect_tmp.write_bytes(final_bytes)
    try:
        doc_for_detect = Document(str(detect_tmp))
        detect_empty_headings_and_order(doc_for_detect, report)
        detect_contrast(doc_for_detect, report)
        with ZipFile(BytesIO(final_bytes), "r") as zf_readonly:
            detect_links(zf_readonly, report)
            detect_tables_merged_empty(zf_readonly, report)
            detect_header_footer(zf_readonly, report)
            # embedded media + gifs (simple)
            media = []
            gifs = []
            for name in zf_readonly.namelist():
                if name.startswith("word/media/") and name.lower().endswith(".gif"):
                    gifs.append(name)
            rels = read_pkg_xml(zf_readonly, "word/_rels/document.xml.rels")
            if rels is not None:
                for rel in rels.findall("Relationship"):
                    t = (rel.get("Type") or "").lower()
                    if "video" in t or "audio" in t:
                        media.append({"id": rel.get("Id"), "target": rel.get("Target"), "type": t})
            report["details"]["embeddedMedia"] = media
            report["details"]["gifsDetected"] = gifs
            report["summary"]["flagged"] += len(media) + len(gifs)
    finally:
        detect_tmp.unlink(missing_ok=True)

    # **Filename suggestion and renaming logic**
    process_file_name(file, report)

    return JSONResponse({
        "fileName": file.filename,
        "suggestedFileName": report["suggestedFileName"],
        "report": report,    
    })

@app.post("/download-document")
async def download_document(file: UploadFile = File(...)):

    if not file:
        raise HTTPException(400, "No file uploaded")
    if not is_docx(file.filename, file.content_type):
        raise HTTPException(400, detail={
            "error": "Please upload a .docx file",
            "details": {"received": {"name": file.filename, "mimetype": file.content_type}},
        })

    # Read the file into memory
    original_bytes = await file.read()

    # Phase A: Apply fixes like table header repeat
    tmp_path = DOWNLOAD_DIR / f"work-{uuid.uuid4().hex}.docx"
    tmp_path.write_bytes(original_bytes)
    doc = Document(str(tmp_path))
    
    # Apply the header repeat fix (table headers, etc.)
    set_table_header_repeat(doc, report={})  # No report needed here
    doc.save(str(tmp_path))
    phase_a_bytes = tmp_path.read_bytes()
    tmp_path.unlink(missing_ok=True)  # Clean up the temp file

    # Phase B: Apply XML replacements (same as upload-document)
    replacements: Dict[str, bytes] = {}

    settings_xml = read_xml_part(phase_a_bytes, "word/settings.xml")
    if settings_xml:
        new_settings = remove_protection_bytes(settings_xml)
        if new_settings:
            replacements["word/settings.xml"] = new_settings

    styles_xml = read_xml_part(phase_a_bytes, "word/styles.xml")
    if styles_xml:
        new_styles = set_default_lang_en_us_bytes(styles_xml)
        if new_styles:
            replacements["word/styles.xml"] = new_styles

    core_xml = read_xml_part(phase_a_bytes, "docProps/core.xml")
    if core_xml:
        new_core = ensure_title_bytes(core_xml)
        if new_core:
            replacements["docProps/core.xml"] = new_core

    # Rebuild the file with all fixes (same logic as upload)
    final_bytes = write_pkg_xml(phase_a_bytes, replacements)

    # **Apply file naming convention** (same as upload-document)
    base_filename = re.sub(r"\.docx$", "", file.filename, flags=re.I)  # Remove the .docx extension
    base_filename = base_filename.replace("_", "-")  # Replace underscores with hyphens
    slugified_filename = slugify(base_filename)  # Apply the slugify function
    suggested_file_name = f"{slugified_filename}.docx"  # Add "-remediated" suffix

    # Validate the rebuilt package before returning it to the client.
    # If validation fails, return a clear JSON error instead of a (possibly corrupt) binary stream.
    import hashlib
    sha256 = hashlib.sha256(final_bytes).hexdigest()

    # Quick OOXML ZIP validation: open as zip and check essential parts exist.
    invalid_reason = None
    try:
        with ZipFile(BytesIO(final_bytes), "r") as zf_check:
            namelist = zf_check.namelist()
            # Minimal required parts for a valid docx
            required = ["[Content_Types].xml", "word/document.xml"]
            missing = [r for r in required if r not in namelist]
            if missing:
                invalid_reason = {"missingParts": missing, "entries": namelist}
    except Exception as e:
        invalid_reason = {"error": str(e)}

    if invalid_reason is not None:
        # Return JSON error with details and a helpful message
        return JSONResponse({
            "error": "remediator_failed",
            "message": "Remediation produced an invalid .docx package",
            "details": invalid_reason,
        }, status_code=500)

    # Now, prepare the remediated file for streaming back to the user and include a SHA256 header
    def iterfile():
        yield final_bytes

    headers = {
        "Content-Disposition": f'attachment; filename="{suggested_file_name}"',
        "X-Docx-SHA256": sha256,
    }

    return StreamingResponse(
        iterfile(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )

# Vercel serverless handler
handler = app
